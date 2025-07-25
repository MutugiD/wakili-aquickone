from langchain_core.tools import tool
from langchain_community.tools.human.tool import HumanInputRun
from typing import Dict, Any, List
from backend.prompts.tools_prompts import (
    demand_letter_system_prompt, demand_letter_user_prompt,
    plaint_system_prompt, plaint_user_prompt,
    document_review_system_prompt, document_review_user_prompt,
    escalation_system_prompt, escalation_user_prompt
)
import json
import os
from backend.document_extraction.extraction_agent import ExtractionAgent
from backend.agent.research_agent.research_agent import ResearchAgent
from backend.agent.memory import AgentMemory
from backend.agent.tooling import ToolRegistry
from langchain_openai import ChatOpenAI

# Human intervention tool for legal review
human_tool = HumanInputRun(
    description="Use this tool when legal document review, validation, or human judgment is required. Use for: document approval, complex legal decisions, client consultation, or when confidence is low."
)

@tool
def draft_demand_letter(request: str) -> str:
    """
    Generate a Kenyan-style demand letter based on client facts and legal precedents.
    Use when a client needs to formally request action from another party (e.g., deposit refund, contract compliance).

    Args:
        request: Detailed description of the client's situation, facts, and what they want. Include any relevant legal precedents.

    Returns:
        Formatted demand letter ready for client review
    """
    # Parse the request to extract components
    try:
        # Simple parsing - in production, use more sophisticated NLP
        components = {
            'client_info': 'To be extracted from request',
            'facts': request,
            'legal_basis': 'Kenyan contract law and relevant statutes',
            'demand': 'Action requested by client',
            'additional_requirements': 'Any specific requirements'
        }

        # Use the structured prompt
        prompt = f"{demand_letter_system_prompt}\n\n{demand_letter_user_prompt.format(**components)}"

        # TODO: Implement LLM call with the structured prompt
        letter_template = f"""
        DEMAND LETTER

        Based on the request: {request}

        [DEMAND LETTER CONTENT TO BE GENERATED USING STRUCTURED PROMPT]

        This document requires human review before sending.
        """
        return letter_template

    except Exception as e:
        return f"Error generating demand letter: {str(e)}. Please escalate to human review."

@tool
def draft_plaint(request: str) -> str:
    """
    Draft a plaint with pleads and grounds for Kenyan courts.
    Use when preparing to file a lawsuit or legal claim.

    Args:
        request: Detailed description of the case facts, legal basis, and relevant precedents

    Returns:
        Download link to the generated plaint document and a source link
    """
    try:
        import os
        from datetime import datetime
        from docx import Document
        # Parse the request to extract components
        components = {
            'case_type': 'Civil case',
            'court': 'Appropriate Kenyan court',
            'plaintiffs': 'To be extracted from request',
            'defendants': 'To be extracted from request',
            'facts': request,
            'causes_of_action': 'Breach of contract, negligence, etc.',
            'legal_authorities': 'Relevant Kenyan statutes and case law',
            'relief_sought': 'Damages, injunctions, etc.',
            'additional_requirements': 'Any specific requirements'
        }

        plaint_text = f"""
        PLAINT\n\nCase Details: {request}\n\n[PLAINT CONTENT TO BE GENERATED USING STRUCTURED PROMPT]\n\nThis document requires human review before filing.
        """
        # Generate Word document
        doc = Document()
        doc.add_heading('PLAINT', 0)
        doc.add_paragraph(f"Case Details: {request}")
        doc.add_paragraph("[PLAINT CONTENT TO BE GENERATED USING STRUCTURED PROMPT]")
        doc.add_paragraph("This document requires human review before filing.")
        # Save to reports directory
        reports_dir = os.path.join(os.getcwd(), 'reports')
        os.makedirs(reports_dir, exist_ok=True)
        filename = f"plaint_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(reports_dir, filename)
        doc.save(file_path)
        # Return download URL and source link (use full backend URL)
        download_url = f"http://localhost:8000/documents/reports/{filename}"
        source_link = "https://www.kenyalaw.org/lex/actview.xql?actid=CAP.%20403"  # Example Traffic Act source
        return f"Draft plaint generated. [Download Word Document]({download_url})\n\nSource: [Kenya Law - Traffic Act (Cap 403)]({source_link})"
    except Exception as e:
        return f"Error generating plaint: {str(e)}. Please escalate to human review."

@tool
def review_legal_document(request: str) -> str:
    """
    Review a legal document for compliance, accuracy, and completeness.
    Use when a document needs professional legal review before finalization.

    Args:
        request: Description of document type and content to review

    Returns:
        Review feedback and recommendations
    """
    try:
        # Parse the request to extract components
        components = {
            'document_type': 'Legal document',
            'purpose': 'To be determined from request',
            'content': request,
            'review_focus': 'Compliance, accuracy, completeness',
            'additional_context': 'Any additional context provided'
        }

        # Use the structured prompt
        prompt = f"{document_review_system_prompt}\n\n{document_review_user_prompt.format(**components)}"

        # TODO: Implement LLM call with the structured prompt
        review_template = f"""
        LEGAL DOCUMENT REVIEW

        Review Request: {request}

        [REVIEW CONTENT TO BE GENERATED USING STRUCTURED PROMPT]

        Status: REQUIRES HUMAN REVIEW
        """
        return review_template

    except Exception as e:
        return f"Error reviewing document: {str(e)}. Please escalate to human review."

@tool
def escalate_to_human(request: str) -> str:
    """
    Escalate the case to human legal professional for review or decision.
    Use when: complex legal issues, high-value cases, unclear facts, or when automated tools cannot handle the request.

    Args:
        request: Description of why human intervention is needed and case context

    Returns:
        Escalation confirmation and next steps
    """
    try:
        # Parse the request to extract components
        components = {
            'case_type': 'Legal matter',
            'parties': 'To be extracted from request',
            'situation': request,
            'issues': 'Complex legal issues requiring human expertise',
            'attempted_actions': 'What has been attempted by automated tools',
            'human_intervention_needed': 'Specific areas requiring human expertise',
            'additional_context': 'Any additional context'
        }

        # Use the structured prompt
        prompt = f"{escalation_system_prompt}\n\n{escalation_user_prompt.format(**components)}"

        # TODO: Implement LLM call with the structured prompt
        escalation_template = f"""
        ESCALATION TO HUMAN LEGAL PROFESSIONAL

        Escalation Request: {request}

        [ESCALATION CONTENT TO BE GENERATED USING STRUCTURED PROMPT]

        Status: ESCALATED
        Next Steps: Human lawyer will review and provide guidance
        """
        return escalation_template

    except Exception as e:
        return f"Error escalating case: {str(e)}. Please contact human lawyer directly."

@tool
def research_kenyan_law(query: str) -> str:
    """
    Research Kenyan case law, statutes, and legal precedents.
    Use when legal research is needed to support document drafting or case analysis.
    Args:
        query: Specific legal research question or topic
    Returns:
        Research results with citations and summaries
    """
    try:
        memory = AgentMemory(memory_type="buffer")
        tools = ToolRegistry()
        agent = ResearchAgent(memory=memory, tool_registry=tools)
        results = agent.handle_request({"type": "search", "query": query})
        result_data = results.get('result')
        # Ensure result_data is a non-empty list of dicts
        if isinstance(result_data, list) and result_data and not all('error' in r for r in result_data):
            return agent.get_research_summary(result_data)
        elif 'error' in results:
            print(f"[research_kenyan_law] Tavily error: {results['error']}")
            return f"Error: {results['error']}"
        else:
            # Fallback to LLM-based answer
            llm = ChatOpenAI(model="gpt-4.1", api_key=os.getenv("OPENAI_API_KEY"), temperature=0)
            prompt = f"You are a Kenyan legal expert. {query} Provide a detailed, step-by-step answer with references to Kenyan law, statutes, and best practices."
            response = llm.invoke(prompt)
            return response.content
    except Exception as e:
        return f"Error in research_kenyan_law: {str(e)}"

@tool
def extract_document(request: str) -> str:
    """
    Extract structured data from a legal document (PDF, DOCX, PNG, JPG, JPEG, etc.).
    Use when a user uploads a document and requests extraction, summary, or analysis.
    Args:
        request: Path to the document file (e.g., 'D:/path/to/file.pdf' or 'docs/outputs/file.png').
    Returns:
        Extracted structured data or summary.
    """
    try:
        import re, os
        # Match all supported file types
        matches = re.findall(r'([A-Za-z]:\\[^\s]+\.(?:pdf|docx|png|jpg|jpeg)|/[^\s]+\.(?:pdf|docx|png|jpg|jpeg))', request, re.IGNORECASE)
        file_path = request.strip()
        if matches:
            file_path = matches[-1]
        # If not absolute, join with outputs dir
        if not os.path.isabs(file_path):
            outputs_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../docs/outputs'))
            file_path = os.path.join(outputs_dir, os.path.basename(file_path))
        print(f"[extract_document] Checking file path: {file_path}")
        agent = ExtractionAgent()
        if os.path.exists(file_path):
            result = agent.extract_structured(file_path)
            return str(result)
        else:
            return f"File not found: {file_path}"
    except Exception as e:
        return f"Error extracting document: {str(e)}"

# Tool registry for the agent
def get_legal_tools() -> List:
    """
    Return all available legal tools including human intervention.
    """
    return [
        draft_demand_letter,
        draft_plaint,
        review_legal_document,
        escalate_to_human,
        research_kenyan_law,
        extract_document,
        human_tool
    ]